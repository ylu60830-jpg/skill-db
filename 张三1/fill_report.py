#!/usr/bin/env python3
"""Fill the 74LS181 experiment report with answers from 内容.txt"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = r"D:\新建1 CLAUDE\张三1\01张三（1）（学生）.docx"
OUTPUT = r"D:\新建1 CLAUDE\张三1\01张三（已完成）.docx"

doc = Document(INPUT)
cell = doc.tables[0].rows[0].cells[0]

# ============================================================
# HELPER: Add a new run to the last run of a paragraph, preserving font
# ============================================================
def add_text_to_paragraph(para, text, bold=None, font_name=None, font_size=None):
    """Add text to existing paragraph by appending a new run with matching formatting."""
    # Strip trailing whitespace from all runs' text elements
    for run in para.runs:
        for t_el in run._r.findall(qn('w:t')):
            if t_el.text:
                t_el.text = t_el.text.rstrip(' ')

    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    # Copy formatting from last run if exists, or use defaults
    if para.runs:
        last_run = para.runs[-1]
        ref = last_run._r.find(qn('w:rPr'))
        if ref is not None:
            rpr = copy.deepcopy(ref)
    if bold is not None:
        b = rpr.find(qn('w:b'))
        if b is None:
            b = OxmlElement('w:b')
            rpr.append(b)
        b.set(qn('w:val'), '1' if bold else '0')
    if font_name:
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        sz = rpr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rpr.append(sz)
        sz.set(qn('w:val'), str(int(font_size / 12700)))  # EMU to half-pt
    run.append(rpr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    para._p.append(run)

def set_paragraph_text(para, text):
    """Replace paragraph text completely, preserving first run formatting."""
    # Clear all runs
    for run in para.runs:
        run._r.getparent().remove(run._r)
    # Add new run with text
    run = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    para._p.append(run)

def fill_paragraph_blanks(para, replacements):
    """Replace blanks in a paragraph. replacements is a list of strings to insert."""
    # This is tricky - we need to find the blank spots.
    # Strategy: rebuild the paragraph text by inserting replacements at "= "
    # For now, we'll use a simpler approach for known patterns
    pass

# ============================================================
# PART 1: Fill paragraph blanks (P61, P62, P65)
# ============================================================

# P61: "令LDDR1=     ，LDDR2=    ，通过KD0-KD7开关输入数据35H..."
# Should become: LDDR1= 1 ，LDDR2= 0
p61 = cell.paragraphs[61]
# Clear existing empty runs after "LDDR1=" and add values
# Strategy: find all text, then rebuild
original_text = p61.text
# Replace multiple spaces with the values
import re
new_text = re.sub(r'LDDR1=\s+，', 'LDDR1= 1 ，', original_text)
new_text = re.sub(r'LDDR2=\s+，', 'LDDR2= 0 ，', new_text)
set_paragraph_text(p61, new_text)
print(f"P61 filled: {new_text[:100]}")

# P62: "令LDDR1=   ，LDDR2=    ，通过KD0-KD7开关输入数据48H..."
p62 = cell.paragraphs[62]
original_text = p62.text
new_text = re.sub(r'LDDR1=\s+，', 'LDDR1= 0 ，', original_text)
new_text = re.sub(r'LDDR2=\s+，', 'LDDR2= 1 ，', new_text)
set_paragraph_text(p62, new_text)
print(f"P62 filled: {new_text[:100]}")

# P65: "关闭数据输入三态门SWB`＝    ，打开ALU输出三态门ALUB`＝    ，当置S3、S2、S1、S0、M为        时，总线指示灯显示     中的数，而置成        时，总线指示灯显示DR2中的数。"
p65 = cell.paragraphs[65]
new_text = "实现过程为：关闭数据输入三态门SWB`＝ 1 ，打开ALU输出三态门ALUB`＝ 0 ，当置S3、S2、S1、S0、M为 11111 时，总线指示灯显示 DR1 中的数，而置成 10101 时，总线指示灯显示DR2中的数。"
set_paragraph_text(p65, new_text)
print(f"P65 filled: {new_text[:100]}")

# ============================================================
# PART 2: Fill Table 2 (DR1=35H, DR2=48H experiment data)
# ============================================================
data_35h_48h = {
    # S3S2S1S0: [M=0,Cn=1, M=0,Cn=0, M=1]
    "0000": ["35H", "36H", "CAH"],
    "0001": ["7DH", "7EH", "82H"],
    "0010": ["B7H", "B8H", "48H"],
    "0011": ["FFH", "00H", "00H"],
    "0100": ["6AH", "6BH", "FFH"],
    "0101": ["B2H", "B3H", "B7H"],
    "0110": ["ECH", "EDH", "7DH"],
    "0111": ["34H", "35H", "35H"],
    "1000": ["35H", "36H", "CAH"],
    "1001": ["7DH", "7EH", "82H"],
    "1010": ["B7H", "B8H", "48H"],
    "1011": ["FFH", "00H", "00H"],
    "1100": ["6AH", "6BH", "FFH"],
    "1101": ["B2H", "B3H", "B7H"],
    "1110": ["B2H", "B3H", "7DH"],
    "1111": ["34H", "35H", "35H"],
}

t2 = cell.tables[2]
# cells[2,3] = M=0,Cn=1; [2,4] = M=0,Cn=0; [2,5] = M=1
for col_idx, val_idx in [(3, 0), (4, 1), (5, 2)]:
    data_cell = t2.rows[2].cells[col_idx]
    for pi, para in enumerate(data_cell.paragraphs):
        if pi < 16:
            # Get the S3S2S1S0 key in order
            s3s2s1s0_list = list(data_35h_48h.keys())
            value = data_35h_48h[s3s2s1s0_list[pi]][val_idx]
            # Add value text after the existing "F＝" content
            add_text_to_paragraph(para, value)
print("Table 2 (35H/48H) filled.")

# ============================================================
# PART 3: Fill Table 3 (学号 DR1=05H, DR2=0BH experiment data)
# ============================================================
data_05h_0bh = {
    "0000": ["05H", "06H", "FAH"],
    "0001": ["0FH", "10H", "F0H"],
    "0010": ["F5H", "F6H", "0AH"],
    "0011": ["FFH", "00H", "00H"],
    "0100": ["09H", "0AH", "FEH"],
    "0101": ["13H", "14H", "F4H"],
    "0110": ["F9H", "FAH", "0EH"],
    "0111": ["03H", "04H", "04H"],
    "1000": ["06H", "07H", "FBH"],
    "1001": ["10H", "11H", "F1H"],
    "1010": ["F6H", "F7H", "0BH"],
    "1011": ["00H", "01H", "01H"],
    "1100": ["0AH", "0BH", "FFH"],
    "1101": ["14H", "15H", "F5H"],
    "1110": ["14H", "15H", "0FH"],
    "1111": ["04H", "05H", "05H"],
}

t3 = cell.tables[3]
# Fill DR1 and DR2 in the first row data cells
set_paragraph_text(t3.rows[2].cells[0].paragraphs[0], "05H")
set_paragraph_text(t3.rows[2].cells[1].paragraphs[0], "0BH")
print("Table 3 (05H/0BH) DR1/DR2 filled.")

# Fill F= values
for col_idx, val_idx in [(3, 0), (4, 1), (5, 2)]:
    data_cell = t3.rows[2].cells[col_idx]
    for pi, para in enumerate(data_cell.paragraphs):
        if pi < 16:
            s3s2s1s0_list = list(data_05h_0bh.keys())
            value = data_05h_0bh[s3s2s1s0_list[pi]][val_idx]
            add_text_to_paragraph(para, value)
print("Table 3 (05H/0BH) data filled.")

# ============================================================
# PART 4: Fill Table 4 (Cy=0 带进位 DR1=8CH, DR2=9FH)
# ============================================================
carry_data_cy0 = [
    # S3S2S1S0, F, Cy, 理论计算
    ("0000", "8DH", "Cy=0", "F=8CH+1=8DH，无进位，Cy=0"),
    ("0001", "A0H", "Cy=0", "F=9FH+1=A0H，无进位，Cy=0"),
    ("0110", "EDH", "Cy=0", "F=8CH-9FH=EDH，有借位（借位标志Cy=0），Cy=0"),
    ("1001", "2CH", "Cy=1", "F=8CH+9FH+1=2CH，有进位，Cy=1"),
    ("1100", "19H", "Cy=1", "F=8CH+8CH+1=19H，有进位，Cy=1"),
    ("1101", "2CH", "Cy=1", "F=9FH+8CH+1=2CH，有进位，Cy=1"),
]

t4 = cell.tables[4]
for col_idx, data_idx in [(3, 1), (4, 2), (5, 3)]:  # F, Cy, theory
    target_cell = t4.rows[1].cells[col_idx]
    # Clear existing paragraph
    for para in target_cell.paragraphs:
        for run in para.runs:
            run._r.getparent().remove(run._r)

    # Add first line to existing paragraph
    para0 = target_cell.paragraphs[0]
    run = OxmlElement('w:r')
    t_el = OxmlElement('w:t')
    t_el.text = carry_data_cy0[0][data_idx]
    t_el.set(qn('xml:space'), 'preserve')
    run.append(t_el)
    para0._p.append(run)

    # Add remaining 5 lines as new paragraphs
    for i in range(1, 6):
        new_p = OxmlElement('w:p')
        new_run = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = carry_data_cy0[i][data_idx]
        new_t.set(qn('xml:space'), 'preserve')
        new_run.append(new_t)
        new_p.append(new_run)
        target_cell._tc.append(new_p)

print("Table 4 (Cy=0) filled.")

# ============================================================
# PART 5: Fill Table 5 (Cy=1 带进位 DR1=8CH, DR2=9FH)
# ============================================================
carry_data_cy1 = [
    # S3S2S1S0, F, Cy, 理论计算 (same values as Cy=0)
    ("0000", "8DH", "Cy=0", "同①，Cy更新为0"),
    ("0001", "A0H", "Cy=0", "同①，Cy更新为0"),
    ("0110", "EDH", "Cy=0", "同①，有借位，Cy变0"),
    ("1001", "2CH", "Cy=1", "同①，有进位，Cy保持1"),
    ("1100", "19H", "Cy=1", "同①，有进位，Cy保持1"),
    ("1101", "2CH", "Cy=1", "同①，有进位，Cy保持1"),
]

t5 = cell.tables[5]
for col_idx, data_idx in [(3, 1), (4, 2), (5, 3)]:
    target_cell = t5.rows[1].cells[col_idx]
    for para in target_cell.paragraphs:
        for run in para.runs:
            run._r.getparent().remove(run._r)

    para0 = target_cell.paragraphs[0]
    run = OxmlElement('w:r')
    t_el = OxmlElement('w:t')
    t_el.text = carry_data_cy1[0][data_idx]
    t_el.set(qn('xml:space'), 'preserve')
    run.append(t_el)
    para0._p.append(run)

    for i in range(1, 6):
        new_p = OxmlElement('w:p')
        new_run = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = carry_data_cy1[i][data_idx]
        new_t.set(qn('xml:space'), 'preserve')
        new_run.append(new_t)
        new_p.append(new_run)
        target_cell._tc.append(new_p)

print("Table 5 (Cy=1) filled.")

# ============================================================
# PART 6: Fill Table 6 (16-bit DR3/DR1=01FEH, DR4/DR2=0102H)
# ============================================================
data_16bit = {
    "0000": ["01FEH", "01FFH", "FE01H"],
    "0001": ["01FEH", "01FFH", "FE01H"],
    "0010": ["FFFFH", "0000H", "0000H"],
    "0011": ["FFFFH", "0000H", "0000H"],
    "0100": ["02FAH", "02FBH", "FEFDH"],
    "0101": ["02FAH", "02FBH", "FEFDH"],
    "0110": ["00FBH", "00FCH", "00FCH"],
    "0111": ["00FBH", "00FCH", "00FCH"],
    "1000": ["0300H", "0301H", "FF03H"],
    "1001": ["0300H", "0301H", "FF03H"],
    "1010": ["0101H", "0102H", "0102H"],
    "1011": ["0101H", "0102H", "0102H"],
    "1100": ["03FCH", "03FDH", "FFFFH"],
    "1101": ["03FCH", "03FDH", "FFFFH"],
    "1110": ["03FCH", "03FDH", "01FEH"],
    "1111": ["01FDH", "01FEH", "01FEH"],
}

t6 = cell.tables[6]
# cells[2,5] = M=0,Cn=1; [2,6] = M=0,Cn=0; [2,7] = M=1
for col_idx, val_idx in [(5, 0), (6, 1), (7, 2)]:
    data_cell = t6.rows[2].cells[col_idx]
    for pi, para in enumerate(data_cell.paragraphs):
        if pi < 16:
            s3s2s1s0_list = list(data_16bit.keys())
            value = data_16bit[s3s2s1s0_list[pi]][val_idx]
            add_text_to_paragraph(para, value)

print("Table 6 (16-bit) filled.")

# ============================================================
# PART 7: Fill paragraph-level answers (思考题, 总结, 设计内容等)
# IMPORTANT: Process in REVERSE index order so insertions don't shift later indices
# ============================================================

def add_answer_after(cell, para_idx, answer_text):
    """Insert a new paragraph with answer text after the given paragraph index."""
    ref_para = cell.paragraphs[para_idx]
    new_p = OxmlElement('w:p')
    # Add paragraph properties matching Normal style
    ppr = OxmlElement('w:pPr')
    pstyle = OxmlElement('w:pStyle')
    pstyle.set(qn('w:val'), 'Normal')
    ppr.append(pstyle)
    new_p.append(ppr)
    # Add run with text
    new_run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), '宋体')
    rpr.append(rfonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # ~12pt
    rpr.append(sz)
    new_run.append(rpr)
    new_t = OxmlElement('w:t')
    new_t.text = answer_text
    new_t.set(qn('xml:space'), 'preserve')
    new_run.append(new_t)
    new_p.append(new_run)
    # Insert after ref_para
    ref_para._p.addnext(new_p)

# Build ALL insertions in one dict, then process in descending index order
all_insertions = {
    # Part 1
    137: "实验一总结：通过手动设置功能开关，完整验证了8位ALU（74LS181）的全部16种算术与逻辑功能。实测结果与理论值吻合，加深了对ALU数据通路、锁存器控制及三态门隔离的理解。",
    140: "答：置数时主要控制LDDR1/LDDR2锁存数据，ALU输出三态门关闭（ALUB'=1），S3～S0、M、Cn原则上可任意。但为避免总线冲突，通常将M=1，S3～S0设为1111（此时ALU输出为A），但因输出三态门关闭，不会影响总线。也可以设为无关项。",
    141: "答：LDDR1为高电平有效锁存控制信号。置数完成后若仍保持高电平（不关闭），锁存器将处于透明状态，总线上的任何数据变化都会立即写入DR1，导致原有数据被意外覆盖。",
    142: "答：74LS181的运算结果经三态门（74LS245）输出至内部总线。ALUB'为低电平有效控制信号，当ALUB'=0时三态门打开，运算结果才能驱动总线，进而通过LED显示灯观察；若ALUB'=1，三态门呈高阻态，总线无法获取ALU的输出。",
    # Part 2
    240: "实验二总结：掌握了带进位控制的算术运算方法，明确了进位标志CY的锁存条件、清零方式及其与CN输入的区别，为多字节运算的进位处理打下基础。",
    243: "答：方法一：令 AR=0，设置 S3~S0、M=00000，使 F=A，DR1 中置入小值（如00H），按动 T4 单脉冲，Cy 即被清零；方法二：将总清开关拨至\"0\"电平直接清零。",
    244: "答：在新的独立运算开始、多字节运算的最低字节、或者需要确定初始进位为0时，必须复位 Cy，防止上次运算的残留进位影响本次结果。",
    245: "答：由图2.1，Cy 锁存在 D 触发器（UN5B）中。其 CLK 受 T4 和 AR 控制。只有当 AR=0（低有效）且 T4 脉冲到来时，CLK 才有上升沿，才能将本次运算的 CN4 状态锁入 CY。条件：AR=0 且有 T4 脉冲。",
    # Part 3 - Design content
    345: "答：利用逻辑与功能。设置 M=1，S3~S0=1011（F＝A·B）。数据 00FFH 存入 DR3/DR1，2D5AH 存入 DR4/DR2，打开 ALUB' 观察输出。",
    347: "答：① 按16位实验接线，查线后上电。② ALUB'=1，SWB'=0，依次将 DR1←FFH，DR2←5AH，DR3←00H，DR4←2DH（通过 LDDR1~4 和 T4 脉冲）。③ 令 SWB'=1，ALUB'=0，LDDR1~4=0。④ 置功能开关 S3~S0=1011，M=1。⑤ 记录总线显示。",
    349: "答：005AH",
    351: "答：00FFH ∧ 2D5AH = 005AH。与实验一致。",
    354: "答：算术加法。M=0，S3~S0=1001（F＝A+B），Cn=1（无进位输入）。",
    356: "答：BBBBH，进位标志 Cy=0。",
    358: "答：1234H + A987H = BBBBH。一致。",
    361: "答：减法。M=0，S3~S0=0110，Cn=0（F＝A－B）。A=7C69H，B=1234H。",
    363: "答：6A35H，借位标志 Cy=1（无借位）。",
    365: "答：两数之差为 6A35H。",
    368: "答：A=1234H，B=7C69H，其余同题3。",
    370: "答：95CBH，借位标志 Cy=0（有借位）。",
    372: "答：95CBH 为差值的补码（即 －6A35H 在16位补码下的表示）。",
    374: "答：题3得 6A35H，Cy=1；题4得 95CBH，Cy=0。两者绝对值相等（6A35H+95CBH=10000H），且借位标志正确指示了减法的借位状态，验证了74LS181减法功能的正确性。",
    376: "实验三总结：将数据通路扩展至16位，成功实现了16位算术逻辑运算，进一步理解了级联间进位传递和16位数据的存取方式。设计题独立设置功能完成了与、加、减等运算，结果正确，有效锻炼了实验设计能力。",
}

# Process in descending index order (highest first) to avoid index shifting
for p_idx in sorted(all_insertions.keys(), reverse=True):
    add_answer_after(cell, p_idx, all_insertions[p_idx])
    print(f"Answer added after P{p_idx}")

print(f"Total paragraph answers inserted: {len(all_insertions)}")

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f"\n{'='*60}")
print(f"DONE! Saved to: {OUTPUT}")
print(f"{'='*60}")
